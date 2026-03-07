#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成Word格式的项目开发更新日志
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_changelog():
    doc = Document()
    
    # 标题
    title = doc.add_heading('怖客 - 项目开发更新日志', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 项目概述
    doc.add_heading('项目概述', level=1)
    doc.add_paragraph(
        '怖客 是一个恐怖灵异内容聚合平台，主要聚合B站UP主"道听途说"系列视频内容，'
        '提供视频浏览、搜索、分类等功能。'
    )
    
    # 开发历程
    doc.add_heading('开发历程', level=1)
    
    # Phase 1
    doc.add_heading('Phase 1: 爬虫系统开发', level=2)
    doc.add_heading('2026-03-06 初期爬虫开发', level=3)
    
    p = doc.add_paragraph()
    p.add_run('创建 yt-dlp 爬虫').bold = True
    p.add_run(' (scripts/ytdlp_crawler.py)')
    
    doc.add_paragraph('• 基于 yt-dlp 库开发B站视频信息爬虫', style='List Bullet')
    doc.add_paragraph('• 支持获取UP主空间所有视频BV号', style='List Bullet')
    doc.add_paragraph('• 支持批量获取视频详细信息', style='List Bullet')
    doc.add_paragraph('• 成功获取UP主 28346875 的 407 个视频', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('解决API认证问题').bold = True
    doc.add_paragraph('• 初期遇到"访问权限不足"错误', style='List Bullet')
    doc.add_paragraph('• 最终采用 yt-dlp 库绕过WBI签名验证', style='List Bullet')
    
    doc.add_heading('2026-03-06 爬虫系统整合', level=3)
    p = doc.add_paragraph()
    p.add_run('创建自动化爬虫编排系统').bold = True
    p.add_run(' (scripts/auto_crawler.py)')
    
    doc.add_paragraph('• 整合三个爬虫程序：BV爬虫、视频信息爬虫、评论爬虫', style='List Bullet')
    doc.add_paragraph('• 实现按顺序依次循环执行的自动化流程', style='List Bullet')
    doc.add_paragraph('• 添加错误处理、日志记录、状态追踪', style='List Bullet')
    doc.add_paragraph('• 支持命令行参数配置', style='List Bullet')
    
    # Phase 2
    doc.add_heading('Phase 2: 数据采集与处理', level=2)
    
    doc.add_heading('2026-03-06 多样化数据采集', level=3)
    p = doc.add_paragraph()
    p.add_run('创建多样化采集器').bold = True
    p.add_run(' (scripts/diverse_collector.py)')
    
    doc.add_paragraph('• 实现5批次采集，共获取15个不同视频的评论区数据', style='List Bullet')
    doc.add_paragraph('• 采用区间分布选择策略确保视频多样性', style='List Bullet')
    doc.add_paragraph('• 总计采集 6400 条评论', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('创建数据处理器').bold = True
    p.add_run(' (scripts/data_processor.py)')
    
    doc.add_paragraph('• 实现 TF-IDF 关键词提取算法', style='List Bullet')
    doc.add_paragraph('• 实现 TextRank 关键词提取算法', style='List Bullet')
    doc.add_paragraph('• 实现混合提取方法（结合两种算法）', style='List Bullet')
    
    doc.add_heading('2026-03-06 CSV数据整合', level=3)
    p = doc.add_paragraph()
    p.add_run('创建CSV整合器').bold = True
    p.add_run(' (scripts/csv_integrator.py)')
    
    doc.add_paragraph('• 处理 Downloads 目录下的 10 个CSV文件', style='List Bullet')
    doc.add_paragraph('• 字段标准化映射', style='List Bullet')
    doc.add_paragraph('• 下载并本地化存储 365 张封面图片', style='List Bullet')
    
    # Phase 3
    doc.add_heading('Phase 3: 前端开发', level=2)
    
    doc.add_heading('2026-03-06 全部视频页面开发', level=3)
    p = doc.add_paragraph()
    p.add_run('创建视频列表页面').bold = True
    p.add_run(' (packages/web/src/app/videos/page.tsx)')
    
    doc.add_paragraph('• 显示视频封面、标题、播放量、评论数', style='List Bullet')
    doc.add_paragraph('• 实现一键跳转B站功能', style='List Bullet')
    doc.add_paragraph('• 实现复制BV号功能', style='List Bullet')
    
    doc.add_heading('2026-03-06 分页功能实现', level=3)
    p = doc.add_paragraph()
    p.add_run('分页系统').bold = True
    
    doc.add_paragraph('• 每页显示 12 个视频（4行×3列布局）', style='List Bullet')
    doc.add_paragraph('• 实现上一页/下一页按钮', style='List Bullet')
    doc.add_paragraph('• 智能页码省略显示', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('排序规则').bold = True
    doc.add_paragraph('• 有标号视频：按期数降序排列', style='List Bullet')
    doc.add_paragraph('• 无标号视频：按发布时间降序排列', style='List Bullet')
    
    # Phase 4
    doc.add_heading('Phase 4: 数据补充与完善', level=2)
    
    doc.add_heading('2026-03-06 期数补充', level=3)
    p = doc.add_paragraph()
    p.add_run('整合新CSV文件').bold = True
    p.add_run(' (scripts/integrate_new_csv.py)')
    
    doc.add_paragraph('• 处理 space (10).csv 文件', style='List Bullet')
    doc.add_paragraph('• 补充 151-169 期视频数据', style='List Bullet')
    doc.add_paragraph('• 期数范围：1-169期（缺失168期）', style='List Bullet')
    
    # Phase 5
    doc.add_heading('Phase 5: 搜索与分类功能', level=2)
    
    doc.add_heading('2026-03-06 搜索功能开发', level=3)
    p = doc.add_paragraph()
    p.add_run('主页搜索功能').bold = True
    
    doc.add_paragraph('• 视频搜索：支持搜索标题、BV号、期数', style='List Bullet')
    doc.add_paragraph('• 实时搜索下拉预览（显示前6个结果）', style='List Bullet')
    doc.add_paragraph('• 故事搜索：支持搜索故事名称、内容、标签', style='List Bullet')
    
    doc.add_heading('2026-03-06 分类功能实现', level=3)
    p = doc.add_paragraph()
    p.add_run('四个分类').bold = True
    
    # 分类表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '分类'
    hdr_cells[1].text = '说明'
    hdr_cells[2].text = '数量'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '全部'
    row_cells[1].text = '所有视频'
    row_cells[2].text = '407'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = '道听途说正集'
    row_cells[1].text = '有期数标号的视频（不含特辑）'
    row_cells[2].text = '动态计算'
    
    row_cells = table.rows[3].cells
    row_cells[0].text = '道听途说特辑'
    row_cells[1].text = '标题包含"特辑"字样'
    row_cells[2].text = '8'
    
    row_cells = table.rows[4].cells
    row_cells[0].text = '都市传说'
    row_cells[1].text = '不包含"道听途说"字样的视频'
    row_cells[2].text = '动态计算'
    
    # Phase 6
    doc.add_heading('Phase 6: 关键词提取系统', level=2)
    
    doc.add_heading('2026-03-06 关键词提取开发', level=3)
    p = doc.add_paragraph()
    p.add_run('创建关键词提取器').bold = True
    p.add_run(' (scripts/extract_keywords_cookie.py)')
    
    doc.add_paragraph('• 基于评论内容使用 jieba 分词提取关键词', style='List Bullet')
    doc.add_paragraph('• 使用Cookie认证绕过API限制', style='List Bullet')
    doc.add_paragraph('• 每个视频至少获取 50 条有效评论', style='List Bullet')
    doc.add_paragraph('• 每个视频提取 5 个关键词', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('过滤规则').bold = True
    doc.add_paragraph('• 过滤人称代词：大佬、小伙伴、阿澈、投稿人等', style='List Bullet')
    doc.add_paragraph('• 过滤停用词：的、了、是、我、你、他等', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('处理结果').bold = True
    doc.add_paragraph('• 成功处理：406 个视频', style='List Bullet')
    doc.add_paragraph('• 失败：1 个（评论数为0）', style='List Bullet')
    
    doc.add_heading('2026-03-06 关键词优化', level=3)
    p = doc.add_paragraph()
    p.add_run('创建优化脚本').bold = True
    p.add_run(' (scripts/optimize_keywords.py)')
    
    doc.add_paragraph('• 分析关键词频率分布', style='List Bullet')
    doc.add_paragraph('• 识别高频低价值词汇', style='List Bullet')
    doc.add_paragraph('• 移除 81 个低价值词汇', style='List Bullet')
    doc.add_paragraph('• 移除 533 次低价值词汇出现', style='List Bullet')
    
    # 技术栈
    doc.add_heading('技术栈', level=1)
    
    doc.add_heading('后端', level=2)
    doc.add_paragraph('• Python 3.x', style='List Bullet')
    doc.add_paragraph('• yt-dlp: B站视频信息爬取', style='List Bullet')
    doc.add_paragraph('• requests: HTTP请求', style='List Bullet')
    doc.add_paragraph('• jieba: 中文分词和关键词提取', style='List Bullet')
    
    doc.add_heading('前端', level=2)
    doc.add_paragraph('• Next.js 14: App Router架构', style='List Bullet')
    doc.add_paragraph('• React: UI组件', style='List Bullet')
    doc.add_paragraph('• TypeScript: 类型安全', style='List Bullet')
    doc.add_paragraph('• Tailwind CSS: 样式框架', style='List Bullet')
    doc.add_paragraph('• pnpm: 包管理器', style='List Bullet')
    
    doc.add_heading('数据存储', level=2)
    doc.add_paragraph('• JSON: 数据库文件格式', style='List Bullet')
    doc.add_paragraph('• 本地文件: 封面图片存储', style='List Bullet')
    
    # 数据统计
    doc.add_heading('数据统计', level=1)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    stats = [
        ('指标', '数值'),
        ('总视频数', '407'),
        ('有期数视频', '168期（1-169，缺168期）'),
        ('无期数视频', '239'),
        ('特辑视频', '8'),
        ('封面图片', '365'),
        ('关键词总数', '~2000'),
        ('评论采集', '6400+'),
    ]
    
    for i, (label, value) in enumerate(stats):
        row_cells = table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
    
    # 项目结构
    doc.add_heading('项目结构', level=1)
    
    structure = """
d:\\怖客\\
├── packages/
│   └── web/                    # Next.js前端项目
│       ├── src/
│       │   └── app/
│       │       ├── page.tsx    # 主页
│       │       └── videos/
│       │           └── page.tsx # 全部视频页面
│       └── public/
│           ├── data/
│           │   └── buke_all_episodes.json  # 视频数据库
│           └── covers/         # 封面图片目录
│
└── scripts/                    # Python脚本目录
    ├── ytdlp_crawler.py        # yt-dlp爬虫
    ├── auto_crawler.py         # 自动化爬虫编排
    ├── diverse_collector.py    # 多样化采集器
    ├── data_processor.py       # 数据处理器
    ├── csv_integrator.py       # CSV整合器
    ├── integrate_new_csv.py    # 新CSV整合
    ├── extract_keywords_cookie.py  # 关键词提取
    └── optimize_keywords.py    # 关键词优化
"""
    
    p = doc.add_paragraph()
    p.add_run(structure).font.name = 'Consolas'
    
    # 版本信息
    doc.add_heading('版本信息', level=1)
    
    p = doc.add_paragraph()
    p.add_run('v1.0.0 - 2026-03-06').bold = True
    doc.add_paragraph('• 初始版本发布', style='List Bullet')
    doc.add_paragraph('• 完成爬虫系统开发', style='List Bullet')
    doc.add_paragraph('• 完成前端基础功能', style='List Bullet')
    doc.add_paragraph('• 完成关键词提取与优化', style='List Bullet')
    
    # 文档信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f'文档生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存文档
    doc.save('怖客_项目开发更新日志.docx')
    print('Word文档已生成: 怖客_项目开发更新日志.docx')


if __name__ == '__main__':
    create_changelog()
